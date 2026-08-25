#!/usr/bin/env python3
"""Convert artifacts/change-log.csv (+ cover from MD stats) to Word .docx."""
from __future__ import annotations

import csv
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, Cm

ROOT = Path(__file__).resolve().parents[2]
CSV_PATH = ROOT / "artifacts" / "change-log.csv"
OUT_DOCX = ROOT / "artifacts" / "change-log.docx"
IST = timezone(timedelta(hours=5, minutes=30))

COLUMNS = [
    "Date time",
    "Platform",
    "Version",
    "What changed",
    "Why",
    "What affected",
    "Criticality",
    "Impact / broke",
    "Overlooked",
]


def sanitize(text: str) -> str:
    return (
        (text or "")
        .replace("\u2014", " - ")
        .replace("\u2013", "-")
        .replace("\u2192", "->")
        .replace("\u2026", "...")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width = new_w
    section.page_height = new_h
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tePr if False else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = sanitize(h)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
        set_cell_shading(hdr[i], "D9E2F3")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = sanitize(val)[:400]
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)
            if ci < len(headers) and headers[ci] == "Criticality" and val == "HIGH":
                set_cell_shading(cells[ci], "F8CBAD")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


def main() -> None:
    if not CSV_PATH.exists():
        raise SystemExit(f"Missing {CSV_PATH} — run generate-full-change-log.py first")

    with CSV_PATH.open(encoding="utf-8", newline="") as f:
        all_rows = list(csv.DictReader(f))

    doc = Document()
    set_landscape(doc.sections[0])

    doc.add_heading(sanitize("Fast Close / wrl-dashboard — Full Change Log"), level=0)
    p = doc.add_paragraph()
    p.add_run(sanitize(f"Generated: {datetime.now(IST).strftime('%Y-%m-%d %H:%M:%S %z')} (IST)")).bold = True

    git_n = sum(1 for r in all_rows if r.get("Platform", "").startswith("Git"))
    ver_n = sum(1 for r in all_rows if r.get("Platform") == "Vercel")
    vps_n = sum(1 for r in all_rows if r.get("Platform") == "VPS")
    high_n = sum(1 for r in all_rows if r.get("Criticality") == "HIGH")
    period = f"{all_rows[0]['Date time']} -> {all_rows[-1]['Date time']}" if all_rows else "n/a"

    for line in [
        f"Period: {period}",
        f"Total rows: {len(all_rows)} (Git(+Vercel) commits: {git_n}, Vercel deploy events: {ver_n}, VPS events: {vps_n})",
        f"HIGH criticality rows: {high_n}",
        "Version scheme: git short SHA (no semver tags; package.json stays 0.1.0)",
        "Sources: local git log, vercel ls --format=json (wrl-dashboard), SSH snapshot of /opt/wrl/database/fast-close-app",
    ]:
        doc.add_paragraph(sanitize(line))

    doc.add_heading("How to read", level=1)
    for line in [
        "Criticality is a heuristic from paths/messages (HIGH = schema/auth/deploy/workers).",
        "Impact / broke only claims breakage when git message or VPS systemd evidence supports it.",
        "Platform Git+Vercel means the commit SHA also appears in a Vercel deployment.",
        "VPS history gap: releases/ keeps only recent SHAs; older VPS flips are not on disk.",
        "CSV/Markdown companions: artifacts/change-log.csv and artifacts/change-log.md",
    ]:
        doc.add_paragraph(sanitize(line), style="List Bullet")

    # High shortlist
    doc.add_heading("High-criticality shortlist", level=1)
    high_headers = ["Date time", "Platform", "Version", "What changed", "Impact / broke", "Overlooked"]
    high_rows = [
        [r.get(h, "") for h in high_headers]
        for r in all_rows
        if r.get("Criticality") == "HIGH"
    ]
    # Cap Word size: if huge, take first/last and note
    note = ""
    if len(high_rows) > 250:
        note = f"Showing first 150 and last 100 of {len(high_rows)} HIGH rows."
        high_rows = high_rows[:150] + high_rows[-100:]
    if note:
        doc.add_paragraph(sanitize(note))
    add_table(doc, high_headers, high_rows, col_widths=[1.3, 0.8, 0.8, 2.2, 2.0, 2.0])

    # Master by month
    doc.add_page_break()
    doc.add_heading("Master log (chronological, by month)", level=1)

    by_month: dict[str, list[dict]] = {}
    for r in all_rows:
        key = (r.get("Date time") or "")[:7] or "unknown"
        by_month.setdefault(key, []).append(r)

    widths = [1.2, 0.75, 0.7, 1.6, 1.2, 1.2, 0.6, 1.3, 1.3]
    for month in sorted(by_month.keys()):
        doc.add_heading(sanitize(month), level=2)
        month_rows = [[r.get(c, "") for c in COLUMNS] for r in by_month[month]]
        # Word can struggle with enormous tables; chunk 80
        chunk = 80
        for i in range(0, len(month_rows), chunk):
            part = month_rows[i : i + chunk]
            if i:
                doc.add_paragraph(sanitize(f"{month} (continued {i + 1}-{i + len(part)})"))
            add_table(doc, COLUMNS, part, col_widths=widths)

    doc.add_heading("Platform notes", level=1)
    doc.add_heading("Git", level=2)
    doc.add_paragraph(sanitize(f"Full local history included ({git_n} commit rows). Branch tried-for-server-cache is experimental."))
    doc.add_heading("Vercel", level=2)
    doc.add_paragraph(
        sanitize(
            "Project wrl-dashboard -> https://wrl-dashboard.vercel.app. "
            "vercel.json deploys only main/master. Deploy list from Vercel CLI retention."
        )
    )
    doc.add_heading("VPS", level=2)
    doc.add_paragraph(
        sanitize(
            "Install base /opt/wrl/database/fast-close-app (current -> SHA release). "
            "Self-hosted Supabase under /opt/supabase. "
            "Live snapshot: artifacts/change-log-raw/vps-snapshot.txt"
        )
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX} ({len(all_rows)} rows, {high_n} HIGH)")


if __name__ == "__main__":
    main()
