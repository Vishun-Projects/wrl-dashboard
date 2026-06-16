#!/usr/bin/env python3
"""Generate WRL VPS hosting runbook as DOCX (+ Markdown backup)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

OUT_DOCX = Path(__file__).resolve().parent / "WRL-VPS-Hosting-Runbook.docx"
OUT_DOCX_ALT = Path(__file__).resolve().parent / "WRL-VPS-Hosting-Runbook-fixed.docx"
OUT_MD = Path(__file__).resolve().parent / "WRL-VPS-Hosting-Runbook.md"


def sanitize(text: str) -> str:
    """ASCII-friendly text — avoids Word 'file is corrupt' on some Windows builds."""
    return (
        text.replace("\u2014", " - ")
        .replace("\u2013", "-")
        .replace("\u2192", "->")
        .replace("\u2026", "...")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(sanitize(text), level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(sanitize(text))
    if bold:
        run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(sanitize(item), style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(sanitize(item), style="List Number")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(sanitize(text))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = sanitize(h)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = sanitize(val)


def build() -> None:
    doc = Document()
    doc.core_properties.title = "WRL VPS Hosting Runbook"
    doc.core_properties.subject = "Supabase cloud to VPS migration"
    doc.core_properties.author = "WRL Portal"
    doc.core_properties.comments = "Regenerate: python scripts/vps-hosting/generate-runbook-docx.py"

    # Title
    title = doc.add_heading(sanitize("WRL Portal - VPS Hosting Runbook"), 0)
    title.alignment = 0
    add_para(
        doc,
        "Step-by-step guide: migrate Supabase Cloud → self-hosted Supabase on VPS, "
        "keep Next.js on Vercel. Use this document for future cutovers or disaster recovery.",
    )
    add_para(doc, "Project: fast-close-app (wrl-dashboard)  |  Last updated: June 2026", bold=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 1. Architecture
    add_heading(doc, "1. Target architecture")
    add_para(
        doc,
        "After migration, the app uses a split stack: auth and REST go through self-hosted Supabase "
        "on the VPS; reports and sync write directly to Postgres on the same VPS.",
    )
    add_bullets(
        doc,
        [
            "VPS (187.127.145.253): Docker Supabase stack + Caddy HTTPS → api.wrl-fsm.cloud",
            "Vercel: Next.js app (wrl-dashboard.vercel.app) — unchanged hosting",
            "Western CRM (westerncrm.com): sync worker reads call data; writes to VPS Postgres",
            "DNS: Hostinger A record api.wrl-fsm.cloud → 187.127.145.253",
        ],
    )
    add_heading(doc, "Data flow", 2)
    add_bullets(
        doc,
        [
            "Login / JWT: Browser → Vercel → https://api.wrl-fsm.cloud (Supabase Auth via Kong)",
            "Profile / app_users: Vercel API → DATABASE_URL (Postgres pooler port 6543)",
            "Reports: Vercel API → calls_latest_hot and related tables on VPS",
            "Sync daemon (local): CRM proxy → transform → VPS Postgres :5432 direct",
        ],
    )

    add_heading(doc, "2. Prerequisites", 1)
    add_heading(doc, "2.1 Before you start", 2)
    add_numbered(
        doc,
        [
            "Ubuntu VPS with root SSH access (ours: root@187.127.145.253)",
            "Domain A record: api.yourdomain.com → VPS public IP",
            "Git Bash on Windows (for deploy scripts) or Linux/macOS shell",
            "Supabase Cloud project still accessible for one final pg_dump",
        ],
    )

    add_heading(doc, "2.2 Secrets to collect from Supabase Cloud", 2)
    add_para(doc, "Dashboard → Settings → API / Database:")
    add_table(
        doc,
        ["Secret", "Where used", "Notes"],
        [
            ["Legacy JWT Secret", ".env.vps-setup JWT_SECRET", "Settings → JWT Keys → Legacy JWT Secret"],
            ["anon key (JWT)", "ANON_KEY, NEXT_PUBLIC_SUPABASE_ANON_KEY", "Must be JWT format, not sb_publishable_*"],
            ["service_role key (JWT)", "SERVICE_ROLE_KEY", "Server-side only; never expose to browser"],
            ["Database password", "POSTGRES_PASSWORD, CLOUD_DB_PASSWORD, DATABASE_URL", "Same password used in cloud pooler URL"],
            ["Project ref", "POOLER_TENANT_ID", "e.g. ddmapuyghfeoyajxbcjh — part of postgres.REF user"],
        ],
    )

    add_heading(doc, "2.3 Local secrets file", 2)
    add_para(doc, "Create .env.vps-setup in repo root (gitignored). Example:")
    add_code(
        doc,
        """JWT_SECRET='your-legacy-jwt-secret'
POSTGRES_PASSWORD='your-db-password'
ANON_KEY='eyJ...anon-jwt...'
SERVICE_ROLE_KEY='eyJ...service-role-jwt...'
CLOUD_DB_PASSWORD='same-as-supabase-db-password'
CLOUD_POOLER_HOST='aws-1-ap-southeast-1.pooler.supabase.com'
VPS_HOST='root@187.127.145.253'""",
    )

    add_heading(doc, "2.4 Developer access and credentials", 2)
    add_para(
        doc,
        "Share this table with developers who need VPS or Studio access. "
        "Fill in passphrase and passwords offline — do not commit real values to git. "
        "Leave the passphrase row blank in the generated doc if you will type it in Word later.",
    )
    add_table(
        doc,
        ["Item", "Value / where to find it"],
        [
            ["VPS SSH", "root@187.127.145.253"],
            ["VPS hostname", "srv1745879 (Hostinger)"],
            ["Public API domain", "https://api.wrl-fsm.cloud"],
            ["Supabase project ref (pooler user)", "ddmapuyghfeoyajxbcjh"],
            ["SSH key (Windows)", r"%USERPROFILE%\.ssh\id_ed25519"],
            ["SSH public key", r"%USERPROFILE%\.ssh\id_ed25519.pub"],
            ["SSH key comment / label", "wrplcrm@gmail.com (or your team email)"],
            ["SSH key passphrase", "(fill in Word — not stored in repo)"],
            ["Supabase Studio username", "supabase"],
            ["Supabase Studio password", "Same as POSTGRES_PASSWORD in .env.vps-setup / DATABASE_URL"],
            ["Postgres pooler (app / Vercel)", "api.wrl-fsm.cloud:6543 — user postgres.ddmapuyghfeoyajxbcjh"],
            ["Postgres direct (sync daemon)", "api.wrl-fsm.cloud:5432 — user postgres (internal)"],
            ["Legacy JWT secret", "Supabase Cloud → Settings → JWT Keys → Legacy JWT Secret"],
            ["anon / service_role keys", "Supabase Cloud → Settings → API (JWT format, not sb_publishable_*)"],
        ],
    )
    add_para(doc, "Passphrase (for your copy in Word only — delete before sharing externally):", bold=True)
    add_para(doc, " ")
    add_para(doc, " ")
    add_para(doc, " ")

    add_heading(doc, "2.5 SSH key setup (Windows, one-time per developer)", 2)
    add_para(doc, "Run in Command Prompt or PowerShell. Creates ~/.ssh/id_ed25519 unless you choose another path.")
    add_code(
        doc,
        """ssh-keygen -t ed25519 -C "wrplcrm@gmail.com"
type %USERPROFILE%\\.ssh\\id_ed25519.pub""",
    )
    add_para(doc, "Install the public key on the VPS (from Git Bash on your PC, not from inside the VPS):")
    add_code(
        doc,
        """# Option A — ssh-copy-id (Git Bash)
ssh-copy-id -i ~/.ssh/id_ed25519.pub root@187.127.145.253

# Option B — manual (paste one line into /root/.ssh/authorized_keys on VPS)
cat ~/.ssh/id_ed25519.pub | ssh root@187.127.145.253 "mkdir -p .ssh && cat >> .ssh/authorized_keys"
""",
    )
    add_para(doc, "Optional — avoid typing passphrase every tunnel session (Git Bash, current session only):")
    add_code(doc, "eval $(ssh-agent -s) && ssh-add ~/.ssh/id_ed25519")
    add_para(doc, "Verify login:")
    add_code(doc, "ssh root@187.127.145.253")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "3. Scripts overview", 1)
    add_para(doc, "All scripts live in scripts/vps-hosting/. Main entry point:")
    add_code(doc, "bash scripts/vps-hosting/deploy-to-vps.sh [setup|repair|migrate|restore|all]")
    add_table(
        doc,
        ["Command", "What it does"],
        [
            ["setup", "Install Docker, clone Supabase, Caddy, UFW; configure .env; start stack"],
            ["repair", "Rebuild .env from .env.example + your secrets; fixes Kong 'name resolution failed'"],
            ["migrate", "pg_dump from cloud on VPS + pg_restore into self-hosted Postgres 17"],
            ["restore", "Re-run pg_restore only (SKIP_DUMP=true) using /root/supabase_cloud.dump"],
            ["all", "setup then migrate (full first-time migration)"],
        ],
    )

    add_heading(doc, "Individual scripts", 2)
    add_bullets(
        doc,
        [
            "setup-supabase.sh — runs ON VPS: Docker, git clone supabase/docker, Caddy, firewall",
            "repair-supabase-env.sh — runs ON VPS: full .env rebuild, PG17 compose, docker compose up",
            "migrate-db-from-cloud.sh — runs on PC: scp restore-on-vps.sh, SSH to VPS",
            "restore-on-vps.sh — runs ON VPS: ensure PG17, pg_dump 17, pg_restore 17, verify counts",
        ],
    )

    add_heading(doc, "4. Step-by-step migration (what we did)", 1)

    add_heading(doc, "Step 1 — DNS", 2)
    add_numbered(
        doc,
        [
            "Hostinger (or your DNS): create A record api.wrl-fsm.cloud → 187.127.145.253",
            "Wait for propagation (minutes to hours)",
            "Verify: nslookup api.wrl-fsm.cloud",
        ],
    )

    add_heading(doc, "Step 2 — Prepare .env.vps-setup", 2)
    add_numbered(
        doc,
        [
            "Copy secrets from Supabase cloud into .env.vps-setup (see section 2.2)",
            "Set VPS_HOST=root@187.127.145.253",
            "Never commit this file to git",
        ],
    )

    add_heading(doc, "Step 3 — Install Supabase on VPS", 2)
    add_para(doc, "From repo root in Git Bash:")
    add_code(doc, "bash scripts/vps-hosting/deploy-to-vps.sh setup")
    add_para(doc, "This installs:")
    add_bullets(
        doc,
        [
            "Docker (via get.docker.com)",
            "Supabase docker stack at /opt/supabase/docker",
            "Caddy reverse proxy: api.wrl-fsm.cloud → localhost:8000 (Kong)",
            "UFW ports 80, 443, 6543",
            "Postgres 17 via docker-compose.pg17.yml",
        ],
    )
    add_para(doc, "Duration: ~10–20 minutes.", bold=True)

    add_heading(doc, "Step 3b — Repair (if health check fails)", 2)
    add_para(
        doc,
        "If curl health returns 'name resolution failed' or Kong errors, a minimal .env wiped "
        "Supabase defaults. Run repair:",
    )
    add_code(doc, "bash scripts/vps-hosting/deploy-to-vps.sh repair")

    add_heading(doc, "Step 4 — Dump cloud DB and restore on VPS", 2)
    add_code(doc, "bash scripts/vps-hosting/deploy-to-vps.sh migrate")
    add_para(doc, "What happens on the VPS:")
    add_numbered(
        doc,
        [
            "Ensures Postgres 17 is running (recreates volumes/db/data if upgrading from PG15)",
            "Runs docker run postgres:17 pg_dump from cloud (direct host first, pooler fallback)",
            "Saves dump to /root/supabase_cloud.dump (~42 MB for our project)",
            "Stops auth/rest/kong temporarily; pg_restore into postgres database",
            "Verifies auth.users count > 0 and public tables exist",
        ],
    )
    add_para(doc, "If dump already exists and you only need restore:")
    add_code(doc, "bash scripts/vps-hosting/deploy-to-vps.sh restore")

    add_heading(doc, "Step 5 — Update application environment", 2)
    add_para(doc, "Local .env.local (and Vercel Production):")
    add_table(
        doc,
        ["Variable", "Value"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "https://api.wrl-fsm.cloud"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Same JWT anon key as cloud"],
            ["SUPABASE_SERVICE_ROLE_KEY", "Same JWT service_role key"],
            [
                "DATABASE_URL",
                "postgresql://postgres.PROJECT_REF:PASSWORD@api.wrl-fsm.cloud:6543/postgres?pgbouncer=true",
            ],
            ["READ_*_FROM", "postgres (all report read flags)"],
            ["SYNC_WORKER_ENABLED", "true"],
        ],
    )
    add_para(
        doc,
        "Important: READ_*_FROM does NOT fix auth 401. You need the four Supabase/DB vars above "
        "plus a successful DB restore (app_users populated).",
        bold=True,
    )
    add_para(doc, "Remove old cloud URLs:")
    add_bullets(
        doc,
        [
            "https://PROJECT_REF.supabase.co",
            "aws-1-ap-southeast-1.pooler.supabase.com (in DATABASE_URL)",
        ],
    )

    add_heading(doc, "Step 6 — Redeploy Vercel", 2)
    add_numbered(
        doc,
        [
            "Vercel → wrl-dashboard → Settings → Environment Variables → save all vars",
            "Deployments → Redeploy Production (required for NEXT_PUBLIC_* bake-in)",
            "Clear browser cookies for vercel.app; log in again",
        ],
    )

    add_heading(doc, "Step 7 — Verify", 2)
    add_para(doc, "Auth health (PowerShell — use curl.exe):")
    add_code(
        doc,
        'curl.exe -s "https://api.wrl-fsm.cloud/auth/v1/health" -H "apikey: YOUR_ANON_KEY"',
    )
    add_para(doc, "Expected: JSON with GoTrue version.")
    add_bullets(
        doc,
        [
            "Local: npm run dev → login → /report loads with data",
            "Vercel: DevTools → GET /api/auth/me returns 200 (not 401/500)",
            "Vercel logs: no exceed_egress_quota (means cloud Supabase is no longer called)",
        ],
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "5. Developer SSH access and Supabase Studio", 1)

    add_heading(doc, "5.1 How Studio is exposed", 2)
    add_bullets(
        doc,
        [
            "Self-hosted Studio is served through Kong on port 8000 inside the VPS.",
            "Port 54323 is the default in upstream Supabase docs but is NOT what we use for browser access.",
            "Public HTTPS: https://api.wrl-fsm.cloud (Kong — needs apikey header for API routes).",
            "Local Studio UI: SSH tunnel from your PC → localhost:8000 (see below).",
        ],
    )

    add_heading(doc, "5.2 Open Studio from Windows (correct way)", 2)
    add_para(
        doc,
        "Run this on your Windows PC in Git Bash or PowerShell. Use -N so SSH only forwards the port "
        "(no remote shell). Keep this terminal open while using Studio.",
    )
    add_code(
        doc,
        """# Terminal 1 — tunnel only (Git Bash or PowerShell)
ssh -N -L 8000:127.0.0.1:8000 root@187.127.145.253

# Browser (same PC)
http://localhost:8000

# Login
#   Username: supabase
#   Password: POSTGRES_PASSWORD (same as in .env.vps-setup / DATABASE_URL)""",
    )
    add_para(doc, "Terminal 2 — run deploy/migrate/sync while tunnel stays open in Terminal 1.")

    add_heading(doc, "5.3 Common SSH tunnel mistakes", 2)
    add_table(
        doc,
        ["Mistake", "Symptom", "Fix"],
        [
            [
                "Tunnel from inside the VPS (nested ssh root@187.127.145.253 while already on VPS)",
                "channel 3: open failed: connect failed: Connection refused",
                "Exit to your PC prompt (C:\\... or PS C:\\...). Run ssh -N -L from there only.",
            ],
            [
                "Wrong local port (54323 instead of 8000)",
                "Connection refused on localhost:54323",
                "Use 8000:127.0.0.1:8000 — Studio is behind Kong on 8000.",
            ],
            [
                "Forgot -N; opened VPS shell and closed window",
                "Tunnel dies when session ends",
                "Use ssh -N -L ... in a dedicated terminal; leave it running.",
            ],
            [
                "Kong / stack not running on VPS",
                "Tunnel OK but browser cannot connect",
                "On VPS: cd /opt/supabase/docker && docker compose ps && curl -s -o /dev/null -w '%{http_code}' http://127.0.0.1:8000",
            ],
            [
                "PowerShell: first connect asks yes/no; typed nothing",
                "Host key verification failed",
                "Type yes once, or use Git Bash where you already accepted the host key.",
            ],
            [
                "Running apt on Windows after disconnect",
                "apt: command not found",
                "apt runs on the VPS only. Reconnect: ssh root@187.127.145.253",
            ],
        ],
    )

    add_heading(doc, "5.4 Verify stack on VPS (SSH session)", 2)
    add_code(
        doc,
        """ssh root@187.127.145.253
cd /opt/supabase/docker
docker compose ps
curl -s -o /dev/null -w '%{http_code}\\n' http://127.0.0.1:8000
# Expect 401 or 404 from Kong — means port 8000 is listening (not connection refused)

# If studio container stopped, restart full stack (not required for normal Studio via Kong):
docker compose up -d""",
    )
    add_para(doc, "Create profiles storage bucket in Studio if avatar uploads fail.")

    add_heading(doc, "5.5 Optional — Postgres direct tunnel (psql / debugging)", 2)
    add_para(doc, "For local psql against VPS Postgres (bypasses pooler). Terminal 1:")
    add_code(
        doc,
        """ssh -N -L 5432:127.0.0.1:5432 root@187.127.145.253

# Terminal 2 — example (password = POSTGRES_PASSWORD)
psql "postgresql://postgres:PASSWORD@127.0.0.1:5432/postgres"
""",
    )

    add_heading(doc, "6. Fresh VPS bootstrap (manual, before deploy script)", 1)
    add_para(
        doc,
        "If the VPS is a new Hostinger/Ubuntu box, these steps were done once before "
        "bash scripts/vps-hosting/deploy-to-vps.sh setup. The setup script also installs Docker and UFW rules.",
    )
    add_code(
        doc,
        """# On VPS as root (after ssh root@187.127.145.253)
apt update && apt upgrade -y
timedatectl set-timezone Asia/Kolkata
timedatectl

# Firewall — allow SSH first, then enable
ufw allow OpenSSH
ufw enable
ufw status verbose

# Swap (recommended for 8 GB RAM during pg_restore)
fallocate -l 4G /swapfile
chmod 600 /swapfile
mkswap /swapfile
swapon /swapfile
echo '/swapfile none swap sw 0 0' >> /etc/fstab
free -h
df -h""",
    )
    add_para(doc, "Then from your PC repo root:")
    add_code(doc, "bash scripts/vps-hosting/deploy-to-vps.sh setup")

    add_heading(doc, "7. Sync worker (CRM → VPS Postgres)", 1)
    add_para(doc, "Run locally to keep hot tables updated while app is open or headless:")
    add_code(doc, "npm run sync-worker:daemon")
    add_para(doc, "Requires in .env.local:")
    add_bullets(
        doc,
        [
            "SYNC_WORKER_ENABLED=true",
            "DATABASE_URL pointing at VPS (daemon uses direct :5432 via USE_DIRECT_DATABASE)",
            "Network must resolve westerncrm.com (CRM) and api.wrl-fsm.cloud (Postgres)",
        ],
    )
    add_para(doc, "After long downtime, run once before daemon:")
    add_code(doc, "npm run sync-worker:incremental")
    add_para(doc, "Incremental fetch uses 7-day CRM chunks to avoid timeouts.")

    add_heading(doc, "8. Troubleshooting", 1)

    issues = [
        (
            "Kong: name resolution failed",
            "Minimal .env missing Supabase service hostnames.",
            "bash scripts/vps-hosting/deploy-to-vps.sh repair",
        ),
        (
            "supabase-db is unhealthy after PG17 switch",
            "docker compose down -v does NOT delete volumes/db/data (PG15 files remain).",
            "bash scripts/vps-hosting/deploy-to-vps.sh restore (or rm -rf volumes/db/data on VPS)",
        ),
        (
            "pg_dump: command not found (Windows)",
            "No local PostgreSQL tools.",
            "Dump runs on VPS via Docker postgres:17 (already in scripts).",
        ),
        (
            "pg_restore unsupported version",
            "PG client version mismatch.",
            "Use docker run postgres:17 pg_restore (script does this).",
        ),
        (
            "Vercel pages stuck on Loading…",
            "401 on /api/auth/me — wrong env or missing app_users.",
            "Set 4 auth/DB vars; redeploy; verify restore; clear cookies.",
        ),
        (
            "The server does not support SSL connections",
            "Vercel forced SSL against VPS pooler without TLS.",
            "Code auto-disables SSL for non-supabase.co hosts; redeploy app.",
        ),
        (
            "Register columns empty (ID, Customer)",
            "Compact distribution API missing full columns.",
            "App now uses register bulk API; redeploy latest code.",
        ),
        (
            "sync-worker ENOTFOUND westerncrm.com",
            "No internet / DNS / VPN.",
            "Fix network; daemon retries every 180s.",
        ),
        (
            "sync-worker CRM timeout",
            "Huge incremental window after downtime.",
            "Chunked fetch + manual incremental run once.",
        ),
        (
            "SSH tunnel: channel 3 connection refused",
            "Ran ssh -L from inside VPS, or wrong port (54323).",
            "On Windows PC: ssh -N -L 8000:127.0.0.1:8000 root@187.127.145.253; open localhost:8000.",
        ),
        (
            "SSH: client_loop send disconnect / Connection reset",
            "Idle tunnel, network drop, or nested SSH session.",
            "Close all VPS sessions; one clean tunnel from PC with -N; re-run ssh-add if passphrase key.",
        ),
        (
            "SSH: Host key verification failed (PowerShell)",
            "First connect prompt skipped.",
            "ssh root@187.127.145.253, type yes; or use Git Bash.",
        ),
    ]
    add_table(doc, ["Symptom", "Cause", "Fix"], list(issues))

    add_heading(doc, "9. Post-migration checklist", 1)
    add_numbered(
        doc,
        [
            "auth.users and app_users have rows on VPS",
            "Login works on localhost and Vercel",
            "/report and /api/auth/me return 200",
            "Sync daemon completes incremental without errors",
            "Create profiles bucket in Studio if needed",
            "Monitor Vercel + VPS for 24–48 hours",
            "Optional: rotate secrets after stable cutover",
            "Optional: decommission Supabase cloud after backup retention period",
        ],
    )

    add_heading(doc, "10. Quick reference — commands", 1)
    add_code(
        doc,
        """# --- Windows: one-time SSH key ---
ssh-keygen -t ed25519 -C "wrplcrm@gmail.com"
type %USERPROFILE%\\.ssh\\id_ed25519.pub
ssh-copy-id -i %USERPROFILE%\\.ssh\\id_ed25519.pub root@187.127.145.253

# --- Studio (Terminal 1 on PC — keep open) ---
ssh -N -L 8000:127.0.0.1:8000 root@187.127.145.253
# Browser: http://localhost:8000  user: supabase  pass: POSTGRES_PASSWORD

# --- Full first-time setup + migrate (Terminal 2, repo root, Git Bash) ---
bash scripts/vps-hosting/deploy-to-vps.sh all

# Fix Kong / .env only
bash scripts/vps-hosting/deploy-to-vps.sh repair

# Re-restore from existing dump
bash scripts/vps-hosting/deploy-to-vps.sh restore

# Health check (PowerShell)
curl.exe -s "https://api.wrl-fsm.cloud/auth/v1/health" -H "apikey: YOUR_ANON_KEY"

# VPS: check Docker stack
ssh root@187.127.145.253 "cd /opt/supabase/docker && docker compose ps"

# Local dev
npm run dev

# Sync daemon
npm run sync-worker:daemon""",
    )

    add_heading(doc, "11. File locations", 1)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["scripts/vps-hosting/", "All migration scripts"],
            [".env.vps-setup", "Local secrets for deploy-to-vps.sh (gitignored)"],
            [".env.local", "Local app config pointing at VPS"],
            ["scripts/vps-hosting/VERCEL_ENV.md", "Vercel cutover checklist"],
            ["scripts/vps-hosting/README.md", "Short README"],
            ["/opt/supabase/docker on VPS", "Supabase stack"],
            ["/root/supabase_cloud.dump on VPS", "Cloud database backup"],
            ["docs/sync.md", "Sync worker documentation"],
        ],
    )

    add_para(doc, "")
    add_para(
        doc,
        "End of runbook. For questions, see scripts/vps-hosting/README.md and VERCEL_ENV.md in the repository.",
        bold=True,
    )

    try:
        doc.save(OUT_DOCX)
        saved = OUT_DOCX
    except PermissionError:
        doc.save(OUT_DOCX_ALT)
        saved = OUT_DOCX_ALT
        print(f"NOTE: {OUT_DOCX.name} is locked (close Word) — wrote {OUT_DOCX_ALT.name} instead")

    # Markdown backup (opens reliably in Word via File > Open)
    OUT_MD.write_text(
        "# WRL Portal - VPS Hosting Runbook\n\n"
        "See WRL-VPS-Hosting-Runbook.docx for formatted version.\n\n"
        "Regenerate both: python scripts/vps-hosting/generate-runbook-docx.py\n\n"
        "Full content is in the DOCX; use scripts/vps-hosting/README.md and VERCEL_ENV.md for quick reference.\n",
        encoding="utf-8",
    )
    print(f"Wrote {saved}")
    print(f"Wrote {OUT_MD}")


if __name__ == "__main__":
    build()
